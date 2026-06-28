#!/usr/bin/env python3
"""
build_docx.py — convert the game design Markdown docs to formatted Word (.docx).

Run:  python build_docx.py
Reads the LIVE .md files and writes game-mechanics.docx, player.docx, mentor.docx
next to them. Designed so the result converts cleanly to Google Docs (uses Word's
built-in Title/Heading styles so the TOC + outline survive the upload).

Formatting (per the user's brief):
  - Times New Roman, 12pt body, throughout.
  - Heading hierarchy (bold, TNR): Title 18 / H1 14 / H2 12 / H3 12.
  - Keep the docs' own section numbers (1, 4.1, ...); indent sub-parts beneath.
  - Cover title block, auto Table of Contents, page numbers; A4.
  - Native Word tables (font shrunk for wide ones); monospace code diagrams.
"""
import re
import datetime
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
TODAY = datetime.date.today().strftime("%d %B %Y")

# Glyphs Times New Roman can't render -> text equivalents (applied to all text).
EMOJI_MAP = {
    "⬇": "(-)",        # down arrow  -> cash out
    "⬆": "(+)",        # up arrow    -> cash in
    "★": "[skill]",    # black star  -> skill-up card
    "❌": "X",          # cross mark  -> "don't"
    "✅": "[ESCAPED]",  # check mark  -> win flag
    "⛔": "(locked)",   # no entry    -> don't-edit
    "\U0001F3AF": "",       # dart
}
# Box-drawing / pointers -> ASCII (only inside code diagrams, for monospace portability).
BOX_MAP = {
    "─": "-", "│": "|", "┬": "+", "┴": "+", "┼": "+",
    "┐": "+", "┘": "+", "└": "+", "┌": "+", "┤": "+",
    "├": "+", "►": ">", "▼": "v", "▲": "^",
}

SUBTITLES = {
    "game-mechanics": "Comprehensive Game Mechanics",
    "player": "Player Guide",
    "mentor": "Mentor / Banker Guide",
}


def clean(text, code=False):
    for k, v in EMOJI_MAP.items():
        text = text.replace(k, v)
    if code:
        for k, v in BOX_MAP.items():
            text = text.replace(k, v)
    return text


# ---- styles -----------------------------------------------------------------
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 12)):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.italic = False


# ---- fields (TOC, page number) ----------------------------------------------
def _field(run, instr, placeholder=""):
    r = run._r
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = placeholder
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, it, sep, t, e):
        r.append(el)


def add_toc(doc):
    h = doc.add_paragraph("Contents"); h.style = doc.styles["Heading 1"]
    p = doc.add_paragraph()
    _field(p.add_run(), 'TOC \\o "1-3" \\h \\z \\u',
           "Right-click here and choose Update Field to build the contents "
           "(Google Docs rebuilds it automatically on import).")


def add_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ").font.name = BODY_FONT
    _field(p.add_run(), "PAGE", "1")


# ---- inline formatting ------------------------------------------------------
INLINE = re.compile(r"\*\*(?P<b>.+?)\*\*|\*(?P<i>.+?)\*|`(?P<c>.+?)`|\[(?P<lt>.+?)\]\((?P<lu>.+?)\)")


def add_runs(paragraph, text, size=None):
    text = clean(text)
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], size=size)
        if m.group("b") is not None:
            _run(paragraph, m.group("b"), bold=True, size=size)
        elif m.group("i") is not None:
            _run(paragraph, m.group("i"), italic=True, size=size)
        elif m.group("c") is not None:
            _run(paragraph, m.group("c"), mono=True, size=size)
        else:  # link -> visible text only
            _run(paragraph, m.group("lt"), size=size)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], size=size)


def _run(paragraph, text, bold=False, italic=False, mono=False, size=None):
    r = paragraph.add_run(text)
    r.font.name = MONO_FONT if mono else BODY_FONT
    r.font.bold = bold
    r.font.italic = italic
    if size:
        r.font.size = Pt(size)
    elif mono:
        r.font.size = Pt(10)
    return r


# ---- tables -----------------------------------------------------------------
def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def add_table(doc, rows):
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]  # skip the |---| separator
    ncols = len(header)
    size = 8 if ncols >= 8 else 9 if ncols >= 6 else 10 if ncols >= 4 else 11
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, cell in enumerate(table.rows[0].cells):
        add_runs(cell.paragraphs[0], header[i] if i < len(header) else "", size=size)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    for r in body:
        cells = table.add_row().cells
        for i in range(ncols):
            add_runs(cells[i].paragraphs[0], r[i] if i < len(r) else "", size=size)
    doc.add_paragraph()  # spacer after table


# ---- main parse -------------------------------------------------------------
def is_block_start(line):
    """True if this line begins a new Markdown block (so it can't be a soft-wrap continuation)."""
    s = line.strip()
    if not s:
        return True
    if s[0] in "#>|":
        return True
    if s.startswith("```"):
        return True
    if re.fullmatch(r"-{3,}", s):
        return True
    if re.match(r"^\s*[-*]\s", line) or re.match(r"^\s*\d+\.\s", line):
        return True
    return False


def gather(lines, i):
    """Join soft-wrapped continuation lines following line i; return (joined_text, next_i)."""
    parts = [lines[i].strip()]
    i += 1
    while i < len(lines) and lines[i].strip() and not is_block_start(lines[i]):
        parts.append(lines[i].strip())
        i += 1
    return " ".join(parts), i


def render(doc, md):
    lines = md.split("\n")
    i = 0
    seen_title = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # cover title = first '# ' line; skip it in the body
        if not seen_title and stripped.startswith("# "):
            seen_title = True
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # fenced code block
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            for j, bl in enumerate(block):
                r = p.add_run(clean(bl, code=True))
                r.font.name = MONO_FONT
                r.font.size = Pt(9)
                if j < len(block) - 1:
                    r.add_break()
            continue

        # table block
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            if len(block) >= 2:
                add_table(doc, block)
            continue

        # horizontal rule -> skip
        if re.fullmatch(r"-{3,}", stripped):
            i += 1
            continue

        # headings
        m = re.match(r"^(#{2,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}[min(level, 4)]
            p = doc.add_paragraph(style=doc.styles[style])
            add_runs(p, m.group(2))
            i += 1
            continue

        # blockquote (join consecutive '>' lines)
        if stripped.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip("> ").rstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            add_runs(p, " ".join(quote))
            for r in p.runs:
                r.font.italic = True
            continue

        # bullet list
        mb = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if mb:
            lvl = len(mb.group(1)) // 2
            content = mb.group(2)
            i += 1
            while i < len(lines) and lines[i].strip() and not is_block_start(lines[i]):
                content += " " + lines[i].strip()
                i += 1
            style = "List Bullet" if lvl == 0 else f"List Bullet {min(lvl + 1, 3)}"
            try:
                p = doc.add_paragraph(style=doc.styles[style])
            except KeyError:
                p = doc.add_paragraph(style=doc.styles["List Bullet"])
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (lvl + 1))
            add_runs(p, content)
            continue

        # numbered list
        mn = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if mn:
            lvl = len(mn.group(1)) // 2
            content = mn.group(2)
            i += 1
            while i < len(lines) and lines[i].strip() and not is_block_start(lines[i]):
                content += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style=doc.styles["List Number"])
            p.paragraph_format.left_indent = Inches(0.5 + 0.25 * lvl)
            add_runs(p, content)
            continue

        # plain paragraph (indented under its heading)
        text, i = gather(lines, i)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        add_runs(p, text)


def build(md_path, docx_path, key):
    with open(md_path, encoding="utf-8") as f:
        md = f.read()
    title = md.split("\n", 1)[0].lstrip("# ").strip()
    title = clean(title)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)  # A4
    sec.left_margin = sec.right_margin = Inches(1)
    setup_styles(doc)

    # cover
    t = doc.add_paragraph(title, style=doc.styles["Title"])
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(SUBTITLES[key]); rs.font.size = Pt(14); rs.font.bold = True
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = d.add_run(TODAY); rd.font.italic = True
    doc.add_page_break()

    add_toc(doc)
    doc.add_page_break()

    render(doc, md)
    add_page_numbers(doc)
    doc.save(docx_path)
    print(f"  wrote {docx_path}")


if __name__ == "__main__":
    import os
    here = os.path.dirname(os.path.abspath(__file__))
    jobs = [("game-mechanics.md", "game-mechanics.docx", "game-mechanics"),
            ("player.md", "player.docx", "player"),
            ("mentor.md", "mentor.docx", "mentor")]
    print("Building .docx files:")
    for src, out, key in jobs:
        build(os.path.join(here, src), os.path.join(here, out), key)
    print("Done.")
