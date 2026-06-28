"""Minimal Markdown -> .docx converter tailored for the jacket sale plan.
Handles: # ## ### headings, **bold**, `code`, bullet/numbered lists (with
one level of nesting), pipe tables, blockquotes, and code fences.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs(paragraph, text):
    """Add text to a paragraph, rendering **bold** and `code` inline."""
    # split on bold and inline-code tokens, keeping delimiters
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(part)


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line))


def parse_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def main(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    in_code = False
    code_buf = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # code fences
        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # blank line
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.match(r"^---+$", stripped):
            i += 1
            continue

        # tables: a row followed by a separator row
        if "|" in line and i + 1 < n and is_table_sep(lines[i + 1]):
            header = parse_row(line)
            i += 2
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(parse_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                add_runs(cell.paragraphs[0], h)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row[: len(header)]):
                    cells[j].paragraphs[0].text = ""
                    add_runs(cells[j].paragraphs[0], val)
            doc.add_paragraph()
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 4))
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run_text = stripped.lstrip(">").strip()
            add_runs(p, run_text)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        # nested bullet (2-space indent)
        m = re.match(r"^(\s+)-\s+(.*)$", line)
        if m and len(m.group(1)) >= 2:
            p = doc.add_paragraph(style="List Bullet 2")
            add_runs(p, m.group(2))
            i += 1
            continue

        # bullet
        m = re.match(r"^-\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, m.group(1))
            i += 1
            continue

        # numbered
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(1))
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(docx_path)
    print("Saved:", docx_path)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
